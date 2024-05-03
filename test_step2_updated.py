import streamlit as st
import json
from docx import Document
from langflow.processing.load import load_flow_from_json
import os
import re
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

from langchain_chroma import Chroma
from langchain_community.document_loaders import Docx2txtLoader
#from langchain_text_splitters import CharacterTextSplitter
#from langchain_text_splitters import TokenTextSplitter
from langchain_text_splitters import SpacyTextSplitter
import spacy
spacy.load('en_core_web_sm')
from langchain.embeddings import OpenAIEmbeddings

from dotenv import load_dotenv
load_dotenv()

openai_api_key = os.getenv('OPENAI_API_KEY')

# Load initial JSON configuration
with open('carlat.json', "r", encoding="utf-8") as f:
    flow_graph = json.load(f)


def update_api_key(data):
    if isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, dict) or isinstance(value, list):
                update_api_key(value)  # Recurse into sub-dicts or lists
            elif key == 'name' and value == 'openai_api_key':
                if 'value' in data:
                    old_value = data['value']  # Capture old value for logging or other purposes
                    data['value'] = openai_api_key  # Update with new API key
                    print(f"API key updated from {old_value} to {openai_api_key}")
    elif isinstance(data, list):
        for item in data:
            update_api_key(item)  # Recurse into elements if it's a list

# Update the API key in the JSON
update_api_key(flow_graph)


def update_json_with_file_path(uploaded_file):
    # Save the uploaded file and update the JSON configuration
    if uploaded_file is not None:
        base_path = os.getcwd()  # Or any directory you want to save files in
        file_path = os.path.join(base_path, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        # Update the JSON node for the PDF loader with the new file path
        for node in flow_graph["data"]["nodes"]:
            if node["data"]["type"] == "UnstructuredWordDocumentLoader":
                node["data"]["node"]["template"]["file_path"]["file_path"] = file_path
        return file_path
    return None

def load_flow(flow_file):
    return load_flow_from_json(flow_file)


def extract_keywords(flow1, custom_prompt=None):
    if custom_prompt is None:
        custom_prompt = 'Return up to 10 most interesting, important, and unique key topics found in this raw Q&A transcript. Display them as a bulleted list. Only return the topics and do not return the descriptions'
    result = flow1(custom_prompt)
    content_output = result['chat_history'][1].content  # Ensure the correct path to content
    return content_output


def get_quotes(flow1, topics):
    result = flow1(f'For each of the extracted topics, find the relevant quotes from the transcript: {topics}. Return the whole quotes as they have appeared in the transcript. The minimum length of the quotes should be atleast 200 words. If there are multiple quotes against a topic, then extract them separately. DO NOT JOIN MULTIPLE QUOTES USING ELLIPSIS AND NO CHANGES SHOULD BE MADE TO THE QUOTES.')
    quotes_output = result['chat_history'][1].content
    return quotes_output

def get_qa_pairs(flow1, quotes):
    result = flow1(f'For each of the extracted quotes, generate 1 to 3 well-worded comprehensive question and answer pairs from those quotes: {quotes}. Do not mention the name of the interviewers or interviewee in the question answer pairs. The answers should be of minimum 200 words for each of the questions generated from the quotes. Do not miss question-answer pair of any topic.')
    qa_pairs_output = result['chat_history'][1].content
    #print(qa_pairs_output)
    return qa_pairs_output

def highlight_text(doc_path, phrases):
    doc = Document(doc_path)
    found_phrases = set()  # Keep track of phrases that have already been highlighted

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue  # Skip empty paragraphs

        # Normalize the text for case-insensitive searching
        text_lower = text.lower()
        matches = []

        # Find the first occurrence of each phrase in the document
        for phrase in phrases:
            if phrase not in found_phrases:  # Only look for phrases that haven't been highlighted yet
                phrase_lower = phrase.lower()
                start = text_lower.find(phrase_lower)
                if start != -1:
                    end = start + len(phrase)
                    matches.append((start, end))
                    found_phrases.add(phrase)  # Mark this phrase as found

        # Sort matches by start position to ensure they are processed in order
        matches.sort()

        # Clear the existing paragraph and rebuild it with highlights
        para.clear()
        last_end = 0
        for start, end in matches:
            # Add text before the matched phrase if there's any
            if start > last_end:
                para.add_run(text[last_end:start])
            # Add the matched phrase with highlighting
            highlighted_run = para.add_run(text[start:end])
            highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            highlighted_run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure text color is black
            last_end = end

        # Ensure any text after the last match is also added
        if last_end < len(text):
            para.add_run(text[last_end:])
            
    doc.save(doc_path)

def create_docx(content, filename):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filename)

def reset_app():
    # Clear session state or specific keys
    st.session_state.clear()

# Streamlit UI components
st.title("Dr. Carlat's Q/A Editor")

# Reset Button
if st.button("Reset App"):
    reset_app()
    st.experimental_rerun()

embeddings = OpenAIEmbeddings()




uploaded_file = st.file_uploader("Upload a document", type=["docx"])
if uploaded_file:
    file_path = update_json_with_file_path(uploaded_file)
    st.success(f"File uploaded at {file_path}")
    if uploaded_file is not None:
        base_path = os.getcwd()  # Or any directory you want to save files in
        file_path = os.path.join(base_path, uploaded_file.name)
    loader = Docx2txtLoader(file_path)
    documents = loader.load()
    #text_splitter = TokenTextSplitter(chunk_size=256, chunk_overlap=0)
    text_splitter = SpacyTextSplitter(chunk_size=1000)
    docs = text_splitter.split_documents(documents)



    # load it into Chroma
    db = Chroma.from_documents(docs, embeddings)

    flow1 = load_flow(flow_graph)

    use_custom_prompt = st.checkbox('Use custom prompt for extracting keywords')
    custom_prompt = ""
    if use_custom_prompt:
        custom_prompt = st.text_area('Enter your custom prompt for keyword extraction', height=150)


    if st.button("Extract Topics"):
        if use_custom_prompt and custom_prompt:
            keywords = extract_keywords(flow1, custom_prompt=custom_prompt)
        else:
            keywords = extract_keywords(flow1)
        st.session_state.keywords = keywords  # Save keywords to session state

    if "keywords" in st.session_state:
        # Display keywords in an editable text box
        st.session_state.edited_keywords = st.text_area("Edit Keywords", value=st.session_state.keywords, height=300)


    if st.button("Get Quotes"):
        if "keywords" in st.session_state:
            #quotes = get_quotes(flow1, st.session_state.edited_keywords)
            lines = st.session_state.edited_keywords.strip().split('\n')

            # Remove hyphens and strip extra whitespace
            headings = [line.strip('- ').strip() for line in lines]

            formatted_output = []
            for heading in headings:
                docs = db.similarity_search(heading)
                formatted_string = f"- {heading}\n\"{docs[0].page_content}\"\n"
                formatted_output.append(formatted_string)
            quotes = "\n".join(formatted_output)

            st.session_state.quotes = quotes  # Save quotes to session state
        # elif st.session_state.edited_keywords:
        #     quotes = get_quotes(flow1, st.session_state.edited_keywords)
        #     st.session_state.quotes = quotes
        else:
            st.error("Extract keywords first before getting quotes.")

    if st.button("Highlight and Download Document"):
        if "quotes" in st.session_state and uploaded_file:
            # The file_path should have been set when the file was uploaded
            text = st.session_state.quotes
            pattern = r'"([^"]*)"'
            quotes = re.findall(pattern, text)
            quotes = [line.strip() for item in quotes for line in item.split('\n') if line.strip()]
            highlight_text(file_path, quotes)  # Assuming quotes are newline-separated
            st.success("Document highlighted successfully.")
            # Provide a download button for the highlighted document
            with open(file_path, "rb") as file:
                st.download_button(
                    label="Download Highlighted Document",
                    data=file,
                    file_name="highlighted_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Please upload a document and extract quotes before highlighting.")

    if "quotes" in st.session_state:
        # Display quotes in a new text box
        st.text_area("Extracted Quotes", value=st.session_state.quotes, height=300, key="quotes_box")

    if st.button("Get QA pairs"):
        if "quotes" in st.session_state:
            qa_pairs = get_qa_pairs(flow1, st.session_state.quotes)
            st.session_state.qa_pairs = qa_pairs  # Save QA pairs to session state
        else:
            st.error("Generate quotes first before getting QA pairs.")

    if "qa_pairs" in st.session_state:
        # Display QA pairs in a new text box
        st.text_area("Generated QA Pairs", value=st.session_state.qa_pairs, height=300, key="qa_pairs_box")

        if st.button("Download QA Pairs"):
            create_docx(st.session_state.qa_pairs, "qa_pairs.docx")
            with open("qa_pairs.docx", "rb") as file:
                st.download_button(
                    label="Download QA Pairs as DOCX",
                    data=file,
                    file_name="qa_pairs.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    if st.button("Download Keywords"):
        create_docx(st.session_state.edited_keywords if 'keywords' in st.session_state else '', 'output.docx')
        with open("output.docx", "rb") as file:
            st.download_button(
                label="Download Keywords as DOCX",
                data=file,
                file_name="output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
